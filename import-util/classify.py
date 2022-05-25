# standard lib imports
import os
import shutil
import platform

# Third Party lib imports
import boto3
import docx
import patoolib


if os.path.isfile('ExamQuestions.zip'):
    patoolib.extract_archive("ExamQuestions.zip", outdir=".")
else:
    patoolib.extract_archive("ExamQuestions.rar", outdir=".")


path = '{0}/ExamQuestions'.format(os.getcwd())
delimiter = ''

if platform.system() == 'Windows':
    global delimiter
    delimiter = '\\'
else:
    global delimiter
    delimiter = '/'


# Dynamo DB connection
dynamodb = boto3.resource('dynamodb', region_name='us-west-2', endpoint_url="http://localhost:8000")

# Dynamo DB Tables
boards_table = dynamodb.Table('boards')
levels_table = dynamodb.Table('levels')
subjects_table = dynamodb.Table('subjects')
questions_table = dynamodb.Table('questions')
question_types_table = dynamodb.Table('question_types')

# CONSTANTS: All documents in Tables
BOARDS = boards_table.scan()["Items"]
LEVELS = levels_table.scan()["Items"]
SUBJECTS = subjects_table.scan()["Items"]
QUESTION_TYPES = question_types_table.scan()["Items"]
QUESTIONS = questions_table.scan()["Items"]

# global variable
files = []
files_dictionary = {}
all_question_paths = [
    {
        "path": q["path"],
        "question_id": q["question_id"]
    } for q in QUESTIONS]
total_questions_in_DB = len(QUESTIONS)


def get_board_id(board_type):
    for board in BOARDS:
        if board_type == board["name"]:
            return board["board_id"]
    return ''
    pass


def get_level_id(param):
    for level in LEVELS:
        if param == level["level_name"]:
            return level["level_id"]
    return ''
    pass


def get_question_type_id(q_type):
    for question in QUESTION_TYPES:
        if q_type == question["type"]:
            return question["question_type_id"]
    return ''
    pass


def get_subject_id(subject_name):
    for subject in SUBJECTS:
        if subject_name == subject["subject_name"]:
            return subject["subject_id"]
    return ''
    pass


def get_question_path(board_id, level_id, subject_id, chapter_id, question_type_id, question_no):
    if delimiter == '/':
        return get_question_path_linux(board_id, level_id, subject_id, chapter_id, question_type_id, question_no)
    else:
        return get_question_path_windows(board_id, level_id, subject_id, chapter_id, question_type_id, question_no)
    pass


def get_question_path_linux(board_id, level_id, subject_id, chapter_id, question_type_id, question_no):
    return "{}/{}/{}/{}/{}/{}".format(
        board_id,
        level_id,
        subject_id,
        chapter_id,
        question_type_id,
        question_no)
    pass


def get_question_path_windows(board_id, level_id, subject_id, chapter_id, question_type_id, question_no):
    return "{}\\{}\\{}\\{}\\{}\\{}".format(
        board_id,
        level_id,
        subject_id,
        chapter_id,
        question_type_id,
        question_no)
    pass


def get_question_title(question_doc_path):
    question_title = ''
    doc = docx.Document(question_doc_path)
    for i in range(len(doc.paragraphs)):
        if len(question_title) > 97:
            break
        question_title = question_title + doc.paragraphs[i].text + ' '
    return question_title
    pass


# r=root, d=directories, f = files
for r, d, f in os.walk(path):
    marks = []
    board_id = ''
    level_id = ''
    subject_id = ''
    chapter_id = ''
    question_no = ''
    question_type_id = ''

    for file in f:

        # Populate marks array
        if len(marks) == 0:

            no_of_files_current_folder = len(f)
            marks_txt = open(r + delimiter + 'marks.txt', 'r')

            separated_path = r.split(delimiter)

            board_type = separated_path[-5]
            question_type = separated_path[-1]

            board_id = get_board_id(board_type)
            level_id = get_level_id(separated_path[-4])
            subject_id = get_subject_id(separated_path[-3])
            question_type_id = get_question_type_id(question_type)
            chapter_id = separated_path[-2].split('C')[1]

            for i in range(no_of_files_current_folder-1):
                marks.append(int(marks_txt.readline().split('\n')[0]))

        if '.doc' in file:

            # Get complete path of the file
            complete_path = os.path.join(r, file)
            separated_path = complete_path.split(delimiter)

            # if board, level, or Question Type is different than skip this file
            if board_id == '' or level_id == '' or question_type_id == '':
                continue

            question_no = file.split('.')[0]
            qMarks = marks[int(question_no)-1]
            title = get_question_title(complete_path)

            # construct Question path to store in DB
            question_path = get_question_path(board_id, level_id, subject_id, chapter_id, question_type_id, question_no)

            # Insert question in db
            if question_path not in [q["path"] for q in all_question_paths]:
                total_questions_in_DB += 1
                questions_table.put_item(
                    Item={
                        "question_id": total_questions_in_DB,
                        "board_id": board_id,
                        "level_id": level_id,
                        "subject_id": subject_id,
                        "chapter_id": chapter_id,
                        "question_type_id": question_type_id,
                        "title": title,
                        "path": question_path,
                        "marks": qMarks,
                    }
                )
                pass
            else:
                question_id = ''
                for q in all_question_paths:
                    if question_path == q["path"]:
                        question_id = q["question_id"]

                questions_table.update_item(
                    Key={
                        "question_id": question_id
                    },
                    UpdateExpression="set title = :newTitle; set marks = :qMarks",
                    ExpressionAttributeValues={
                        ':newTitle': title,
                        ':qMarks': qMarks,
                    },
                    ReturnValues="UPDATED_NEW"
                )
                pass

            files.append(complete_path)

# delete the extracted rar folder
shutil.rmtree(path)
